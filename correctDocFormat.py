# correctDocFormat.py

import re,os
from docx import Document

# 定义替换函数
def contentFormat(filePath, saveDir=""):
    """
       Word文件格式化。
       参数:
       * file_path -- 文件路径。
       saveDir -- 格式化后的保存路径。
    """
    doc = Document(filePath)
    # 替换每个段落中的手动换行符
    count = 1;
    new_text = [];
    for paragraph in doc.paragraphs:
        lines = re.sub(r'\n+', '\n', paragraph.text).split('\n');
        # print(f"{count}:{lines}");
        count = count +1;
        paragraph.text = "";
        for line in lines:
            doc.add_paragraph(line);
    #若定义了保存路径则保存至新路径，否则直接覆盖原文件
    if(saveDir):
        normalized_path = os.path.normpath(filePath)
        title = os.path.basename(normalized_path);
        filePath = os.path.join(saveDir, f"_{title}");
    doc.save(filePath)


# if __name__ == '__main__':
#     contentFormat(r'J:\storyFile\我养大的妹妹们，只想把我送监狱\改写版本\改_我养大的妹妹们，只想把我送监狱_1-3.docx');
