# downContent2Word.py
# -*- coding: utf-8 -*-
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from selenium import webdriver
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor
import pyperclip,time,os


def toWeb(url, save_dir):
    """
          下载文案素材
          参数:
          * url -- 小说文章网址。
          * saveDir -- 重写后的保存路径。
    """
    driver = webdriver.Chrome()
    driver.implicitly_wait(10)
    driver.get(url)  # 打开百度浏览器
    navList = driver.find_elements_by_class_name("catalogue__item-uUVnjg");
    copyBtn = driver.find_element_by_class_name("novel-kol-icon-file");
    title = driver.find_element_by_class_name("page-title-EwIcIf").text;

    # 创建一个Word文档对象
    # doc = Document()
    # 创建一个新样式并设置字体为微软雅黑
    # style = doc.styles.add_style('MicrosoftYaHei', WD_STYLE_TYPE.PARAGRAPH)
    # style.font.name = '微软雅黑'
    # style.font.size = Pt(12)  # 设置字体大小为12点
    # doc.styles['Normal'].font.name = u'宋体'
    # doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # doc.styles['Normal'].font.size = Pt(10.5)
    # doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
    # save_dir = f'J:\storyFile\{title}'
    save_dir = os.path.join(save_dir, title);
    # 检查目录是否存在
    if not os.path.exists(save_dir):
        # 如果目录不存在，则创建目录
        os.makedirs(save_dir)
        print("目录已创建:", save_dir)
    else:
        # 目录已存在
        print("目录已存在:", save_dir)
    start = 0;
    count = 0;
    isSave = False;
    for nav in navList:
        classname = nav.get_attribute("class")
        if "catalogue__item--disable-AXZUkF" in classname:
            break;
        if(count == start):
            doc = Document();
            doc.styles['Normal'].font.name = u'宋体'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            doc.styles['Normal'].font.size = Pt(10.5)
            doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
        nav.click();
        #执行复制文案操作
        time.sleep(1.5);
        copyBtn.click();
        # 添加一些文本
        parse_text = pyperclip.paste();
        lines = parse_text.split('\n')
        lines.pop(0)
        # 使用str.join('\n')将列表重新组合成字符串
        parse_text = '\n'.join(lines)
        doc.add_paragraph(parse_text);
        count = count + 1;
        if (count-start) == 3:
            file_name = f"{title}_{start+1}-{count}.docx";
            save_path = os.path.join(save_dir, file_name)
            print(f"开始保存文件{file_name}")
            # 保存文档到指定路径
            doc.save(save_path)
            print(f"文件{file_name}保存完成")
            isSave = True;
            start = count;
        else:
            isSave = False;
    driver.quit();
    # 使用os.path模块构建路径
    if not isSave :
        file_name = f"{title}_{start+1}-{count}.docx";
        save_path = os.path.join(save_dir, file_name)
        print(f"开始保存文件{file_name}")
        # 保存文档到指定路径
        doc.save(save_path)
        print(f"文件{file_name}保存完成")
    os.system("start " + save_dir);
    # driver.quit()  # 关闭浏览器


# if __name__ == '__main__':
#     toWeb("https://promoter.fanqieopen.com/page/share/book-detail?token=2fac8af6da4347b1b39a8359414f6719&tab_type=2&key=6&genre=0&book_id=7207072067127086118",
#     r'J:\storyFile');