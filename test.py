import re
from docx import Document


# 定义替换函数
def replace_manual_line_breaks(doc):
    # 替换每个段落中的手动换行符
    count = 1;
    new_text = [];
    for paragraph in doc.paragraphs:
        lines = re.sub(r'\n+', '\n', paragraph.text).split('\n');
        # print(f"{count}:{lines}");
        count = count +1;
        paragraph.text = "";
        for line in lines:
            doc.add_paragraph(line);
        # cp = "\n".join(paragraph.text);
        # paragraph.text = "";

        # for line in cp:
        #     doc.add_paragraph(line.replace("\n",""));
    # 将实际的换行符替换为段落标记


# if __name__ == '__main__':
#     # 创建一个新的文档或加载一个现有的文档
#     doc = Document(r'J:\storyFile\我养大的妹妹们，只想把我送监狱\改写版本\改_我养大的妹妹们，只想把我送监狱_1-3.docx')
#     replace_manual_line_breaks(doc);
#
#
#     # 保存新的Word文档
#     doc.save('output_word_file.docx')
