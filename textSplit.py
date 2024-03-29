from docx import Document
import os
from articleReWriter import sendContent2Rewriter
from startWindows import startInput
from correctDocFormat import contentFormat
def split_text(arr, docx_path, output_dir, txt_prefix, isSave=True, max_chars=1000):
    doc = Document(docx_path)
    total_chars = 0
    txt_count = 1
    current_text = []

    for paragraph in doc.paragraphs:
        paragraph_text = paragraph.text.strip()
        paragraph_chars = len(paragraph_text)

        # 拆分段落成句子
        sentences = paragraph_text.split('。')
        for sentence in sentences:
            sentence_chars = len(sentence)

            # 如果添加当前句子后超过最大字符数，则写入当前文本文件
            if total_chars + sentence_chars > max_chars:
                write_text(arr, output_dir, txt_prefix, txt_count, current_text, isSave)
                txt_count += 1
                current_text = []
                total_chars = 0

            current_text.append(sentence)
            total_chars += sentence_chars

    # 写入最后一个文本文件
    write_text(arr, output_dir, txt_prefix, txt_count, current_text, isSave)

def write_text(arr, output_dir, txt_prefix, txt_count, text, isSave=True):
    if text:
        text_content = "\n".join(text)
        arr.append(text_content);
        if isSave:
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
            with open(os.path.join(output_dir, f"{txt_prefix}_{txt_count}.txt"), "w", encoding="utf-8") as f:
                    f.write(text_content)

def write_to_word_file(file_name, file_path, content):
    if not os.path.exists(file_path):
        os.mkdir(file_path);
    # 检查文件是否存在
    if os.path.exists(os.path.join(file_path, file_name)):
        # 如果文件存在，则打开文件并在文件末尾追加内容
        doc = Document(os.path.join(file_path, file_name))
        doc.add_paragraph(content)
        doc.save(os.path.join(file_path, file_name))
        print("Content added to existing file:", file_name)
    else:
        # 如果文件不存在，则创建新的 Word 文件并写入内容
        doc = Document()
        doc.add_paragraph(content)
        doc.save(os.path.join(file_path, file_name))
        print("New file created and content added:", file_name)

def chatgptReWriter(filePath = r"J:\storyFile\我养大的妹妹们，只想把我送监狱\我养大的妹妹们，只想把我送监狱_1-3.docx", saveDir=""):
    """
       Word文档使用ChatGPT重写。
       参数:
       * file_path -- 文件路径。
       saveDir -- 重写后的保存路径。
    """
    arr = []
    file_path = os.path.dirname(filePath);
    file_name = os.path.basename(filePath);
    # file_path = r'J:\storyFie\我养大的妹妹们，只想把我送监狱'  # 替换为你的文章文件路径
    # file_name = "我养大的妹妹们，只想把我送监狱_1-3.docx"
    if(saveDir):
        end_file_path = saveDir;
    else:
        end_file_path = os.path.join(file_path, "改写版本");

    if(os.path.exists(os.path.join(end_file_path,file_name))):
        print("文件已生成，请勿重复生成！");
        return;
    # return;
    split_text(arr, os.path.join(file_path, file_name), "output_directory", "output_file_prefix", False)
    count = 1
    reWriterContent = [];
    for content in arr:
        print(f"正在改写内容{count}")
        reWriterContent.append(sendContent2Rewriter(content));
        # print(reWriterContent);
        print(f"内容{count}改写完毕！已加入缓存中！")
        count = count + 1;
    file_name = f"改_{file_name}"
    write_to_word_file(file_name, end_file_path, "\n".join(reWriterContent));
    #统一格式，避免出现手动换行符的情况
    contentFormat(os.path.join(end_file_path, file_name));
    os.start(f'explorer "{end_file_path}"')

if __name__ == "__main__":
    chatgptReWriter(filePath=r"J:\storyFile\学姐别怕，我来保护你\学姐别怕，我来保护你_4-6.docx");
#    startInput(chatgptReWriter);